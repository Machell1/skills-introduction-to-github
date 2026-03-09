"""
FNID Report Processing Bot

Extracts structured data from uploaded reports (PDF, DOCX, plain text),
determines the target unit, and generates operational guidance.
"""

import os
import re
from datetime import datetime, timedelta

from .constants import (
    ALL_OFFENCES, ALL_PARISHES, BAIL_STATUS, CASE_CLASSIFICATIONS,
    DRUG_TYPES, DRUG_UNITS, FIREARM_CALIBRES, FIREARM_TYPES,
    INTEL_PRIORITIES, INTEL_SOURCES, OPERATION_TYPES, SEIZURE_LOCATIONS,
    TRIAGE_DECISIONS, UNIT_PORTALS, WARRANT_BASIS,
)


# ---------------------------------------------------------------------------
# Text extraction from files
# ---------------------------------------------------------------------------

def extract_text(filepath):
    """Extract raw text from a PDF, DOCX, or TXT file."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return _extract_docx(filepath)
    elif ext == ".txt":
        with open(filepath, "r", errors="replace") as f:
            return f.read()
    elif ext in (".jpg", ".jpeg", ".png"):
        # Image files need OCR - return placeholder guidance
        return ""
    else:
        return ""


def _extract_pdf(filepath):
    """Extract text from PDF using available libraries."""
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except ImportError:
        pass

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(filepath)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except ImportError:
        pass

    return ""


def _extract_docx(filepath):
    """Extract text from DOCX."""
    try:
        import docx
        doc = docx.Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    except ImportError:
        pass
    return ""


# ---------------------------------------------------------------------------
# Unit detection - which unit(s) should this report go to?
# ---------------------------------------------------------------------------

_UNIT_KEYWORDS = {
    "intel": [
        "intelligence", "intel report", "informant", "source", "triage",
        "tip", "surveillance", "crime stop", "311", "nib", "811",
        "threat assessment", "target", "monitoring", "dissemination",
    ],
    "operations": [
        "operation", "search warrant", "raid", "interception", "interdiction",
        "checkpoint", "joint operation", "snap raid", "controlled delivery",
        "team lead", "warrant execution", "cordon", "zoso",
    ],
    "seizures": [
        "firearm", "pistol", "revolver", "rifle", "shotgun", "ammunition",
        "rounds", "magazine", "calibre", "serial number", "ibis", "etrace",
        "narcotics", "cocaine", "ganja", "cannabis", "heroin", "drugs",
        "seizure", "seized", "contraband", "fentanyl", "mdma",
    ],
    "arrests": [
        "arrest", "arrested", "suspect", "detained", "detainee", "bail",
        "remand", "charge", "charged", "caution", "custody", "lock-up",
        "48 hour", "48-hour", "court date", "arraign",
    ],
    "forensics": [
        "exhibit", "chain of custody", "forensic", "lab submission",
        "ballistic", "certificate", "ifslm", "analyst", "dna",
        "evidence", "sealed", "tagged", "specimen",
    ],
    "registry": [
        "case file", "case register", "case management", "dpp",
        "prosecution", "disclosure", "witness statement", "sop",
        "compliance", "case status", "case reference", "cr form",
    ],
}


def detect_units(text):
    """Determine which unit(s) the report content relates to.

    Returns a list of (unit_key, confidence_score) tuples sorted by score descending.
    """
    text_lower = text.lower()
    scores = {}

    for unit, keywords in _UNIT_KEYWORDS.items():
        score = 0
        for kw in keywords:
            count = text_lower.count(kw)
            if count > 0:
                score += count
        if score > 0:
            scores[unit] = score

    if not scores:
        # Default to intel if no keywords matched
        scores["intel"] = 1

    total = sum(scores.values())
    results = [(u, round(s / total * 100)) for u, s in scores.items()]
    results.sort(key=lambda x: x[1], reverse=True)
    return results


# ---------------------------------------------------------------------------
# Field extraction - pull structured data from report text
# ---------------------------------------------------------------------------

def _find_dates(text):
    """Extract date strings from text."""
    patterns = [
        r"\b(\d{4}-\d{2}-\d{2})\b",
        r"\b(\d{1,2}[/\-]\d{1,2}[/\-]\d{4})\b",
        r"\b(\d{1,2}\s+(?:January|February|March|April|May|June|July|August|"
        r"September|October|November|December)\s+\d{4})\b",
        r"\b(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)"
        r"\s+\d{4})\b",
    ]
    dates = []
    for pat in patterns:
        dates.extend(re.findall(pat, text, re.IGNORECASE))
    return dates


def _find_badge_numbers(text):
    """Extract JCF badge numbers."""
    return re.findall(r"\bJCF-\d{4}\b", text, re.IGNORECASE)


def _find_names(text):
    """Extract potential person names (capitalized sequences)."""
    # Match 2-4 capitalized words in a row (likely names)
    names = re.findall(
        r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b", text
    )
    # Filter out common non-name phrases
    stopwords = {"The", "This", "That", "Case", "Court", "Police", "Officer",
                 "Section", "Report", "Jamaica", "Parish", "Division"}
    return [n for n in names if n.split()[0] not in stopwords]


def _find_serial_numbers(text):
    """Extract potential firearm serial numbers."""
    return re.findall(r"\b([A-Z]{1,4}\d{3,10}[A-Z]?\d*)\b", text)


def _match_options(text, options):
    """Find which predefined options appear in the text."""
    text_lower = text.lower()
    matched = []
    for opt in options:
        if opt.lower() in text_lower:
            matched.append(opt)
    return matched


def _find_quantities(text):
    """Extract quantity + unit pairs."""
    patterns = [
        r"(\d+(?:\.\d+)?)\s*(kg|kilogram|gram|g|lbs?|pounds?|oz|ounces?|"
        r"plants?|tablets?|capsules?|ml|litres?|rounds?)\b",
    ]
    results = []
    for pat in patterns:
        for match in re.finditer(pat, text, re.IGNORECASE):
            results.append((match.group(1), match.group(2)))
    return results


def _find_currency(text):
    """Extract monetary amounts."""
    amounts = re.findall(
        r"(?:JMD|\$|J\$|US\$)\s*([\d,]+(?:\.\d{2})?)", text
    )
    return amounts


def extract_fields(text, target_unit):
    """Extract structured fields from text based on target unit.

    Returns a dict of field_name → extracted_value.
    """
    fields = {}
    today = datetime.now().strftime("%Y-%m-%d")

    dates = _find_dates(text)
    badges = _find_badge_numbers(text)
    names = _find_names(text)
    parishes = _match_options(text, ALL_PARISHES)

    # Common fields
    if dates:
        fields["_dates_found"] = dates
    if badges:
        fields["_badges_found"] = badges
    if names:
        fields["_names_found"] = names[:10]  # Cap at 10
    if parishes:
        fields["_parishes_found"] = parishes

    if target_unit == "intel":
        fields["date_received"] = dates[0] if dates else today
        sources = _match_options(text, INTEL_SOURCES)
        if sources:
            fields["source"] = sources[0]
        priorities = _match_options(text, INTEL_PRIORITIES)
        if priorities:
            fields["priority"] = priorities[0]
        if parishes:
            fields["parish"] = parishes[0]
        if names:
            fields["subject_name"] = names[0]
        fields["substance_of_intel"] = text[:500] if len(text) > 500 else text
        triage = _match_options(text, TRIAGE_DECISIONS)
        if triage:
            fields["triage_decision"] = triage[0]

    elif target_unit == "operations":
        fields["op_date"] = dates[0] if dates else today
        op_types = _match_options(text, OPERATION_TYPES)
        if op_types:
            fields["op_type"] = op_types[0]
        warrants = _match_options(text, WARRANT_BASIS)
        if warrants:
            fields["warrant_basis"] = warrants[0]
        if badges:
            fields["team_lead"] = badges[0]
        # Count firearms/narcotics mentioned
        fa_count = len(re.findall(r"firearm|pistol|rifle|shotgun|gun", text, re.I))
        if fa_count:
            fields["firearms_seized"] = str(fa_count)
        qtys = _find_quantities(text)
        if qtys:
            fields["narcotics_seized"] = f"{qtys[0][0]} {qtys[0][1]}"

    elif target_unit == "seizures":
        fields["seizure_date"] = dates[0] if dates else today
        if parishes:
            fields["parish"] = parishes[0]
        locations = _match_options(text, SEIZURE_LOCATIONS)
        if locations:
            fields["location"] = locations[0]
        fa_types = _match_options(text, FIREARM_TYPES)
        drug_types = _match_options(text, DRUG_TYPES)

        if fa_types:
            fields["_subtype"] = "firearms"
            fields["firearm_type"] = fa_types[0]
            calibres = _match_options(text, FIREARM_CALIBRES)
            if calibres:
                fields["calibre"] = calibres[0]
            serials = _find_serial_numbers(text)
            if serials:
                fields["serial_number"] = serials[0]
            ammo = re.findall(r"(\d+)\s*(?:rounds?|cartridges?|ammunition)", text, re.I)
            if ammo:
                fields["ammo_count"] = ammo[0]
        elif drug_types:
            fields["_subtype"] = "narcotics"
            fields["drug_type"] = drug_types[0]
            qtys = _find_quantities(text)
            if qtys:
                fields["quantity"] = qtys[0][0]
                # Map to standard unit
                unit_map = {"kg": "kg", "kilogram": "kg", "gram": "g",
                            "g": "g", "lb": "lbs", "lbs": "lbs",
                            "pound": "lbs", "pounds": "lbs",
                            "oz": "oz", "ounce": "oz", "ounces": "oz",
                            "plant": "plants", "plants": "plants",
                            "tablet": "tablets", "tablets": "tablets"}
                fields["unit"] = unit_map.get(qtys[0][1].lower(), qtys[0][1])
            amounts = _find_currency(text)
            if amounts:
                fields["est_street_value"] = amounts[0].replace(",", "")
        else:
            fields["_subtype"] = "firearms"  # default

    elif target_unit == "arrests":
        fields["arrest_date"] = dates[0] if dates else today
        if names:
            fields["suspect_name"] = names[0]
        offences = _match_options(text, ALL_OFFENCES)
        if offences:
            fields["offence_1"] = offences[0]
            if len(offences) > 1:
                fields["offence_2"] = offences[1]
        bail = _match_options(text, BAIL_STATUS)
        if bail:
            fields["bail_status"] = bail[0]
        if parishes:
            fields["parish"] = parishes[0]

    elif target_unit == "forensics":
        fields["_subtype"] = "custody"
        if dates:
            fields["received_date"] = dates[0]
        # Check for lab-related content
        if any(kw in text.lower() for kw in ["lab submission", "lab result",
                                               "certificate", "analysis"]):
            fields["_subtype"] = "lab"
        if names:
            fields["seized_by"] = names[0]

    elif target_unit == "registry":
        fields["_subtype"] = "cases"
        if dates:
            fields["registration_date"] = dates[0]
        classifications = _match_options(text, CASE_CLASSIFICATIONS)
        if classifications:
            fields["classification"] = classifications[0]
        if names:
            fields["suspect_name"] = names[0]
        if parishes:
            fields["parish"] = parishes[0]
        offences = _match_options(text, ALL_OFFENCES)
        if offences:
            fields["law_and_section"] = offences[0]

    return fields


# ---------------------------------------------------------------------------
# Guidance generation
# ---------------------------------------------------------------------------

def generate_guidance(text, target_unit, extracted_fields):
    """Generate operational guidance based on the report content.

    Returns a list of guidance dicts: {type, title, message, priority}.
    type: 'action', 'warning', 'info', 'compliance'
    priority: 'high', 'medium', 'low'
    """
    guidance = []
    text_lower = text.lower()

    # --- Cross-unit guidance ---

    # Check if report mentions multiple units
    units_detected = detect_units(text)
    if len(units_detected) > 1 and units_detected[1][1] > 20:
        other = units_detected[1][0]
        guidance.append({
            "type": "info",
            "title": "Cross-Unit Relevance",
            "message": f"This report also has significant relevance to the "
                       f"{UNIT_PORTALS[other]['name']}. Consider sharing "
                       f"or creating a linked record there.",
            "priority": "medium",
        })

    # Check for urgency indicators
    urgency_words = ["urgent", "immediate", "critical", "emergency", "threat to life",
                     "armed", "imminent", "time-sensitive"]
    if any(w in text_lower for w in urgency_words):
        guidance.append({
            "type": "warning",
            "title": "Urgent Matter Detected",
            "message": "This report contains urgency indicators. Ensure the "
                       "DCO/DDI is notified immediately and the record is "
                       "submitted (not left as Draft).",
            "priority": "high",
        })

    # Check for names of known persons across existing records
    # (Would query DB in production; here we note the capability)
    names = extracted_fields.get("_names_found", [])
    if names:
        guidance.append({
            "type": "info",
            "title": "Persons Identified",
            "message": f"The bot identified {len(names)} potential person name(s): "
                       f"{', '.join(names[:3])}. Verify these are correct and "
                       f"check existing records for matches.",
            "priority": "low",
        })

    # --- Unit-specific guidance ---

    if target_unit == "intel":
        if "priority" not in extracted_fields:
            guidance.append({
                "type": "action",
                "title": "Priority Not Detected",
                "message": "The bot could not determine the intelligence priority "
                           "level. Please set the priority (Critical/High/Medium/Low) "
                           "before submitting.",
                "priority": "medium",
            })
        if extracted_fields.get("priority") == "Critical":
            guidance.append({
                "type": "warning",
                "title": "Critical Intelligence",
                "message": "This intel is classified as Critical. Per SOP, the DDI "
                           "must be briefed within 1 hour. Triage decision and "
                           "action plan are required immediately.",
                "priority": "high",
            })
        if "triage_decision" not in extracted_fields:
            guidance.append({
                "type": "action",
                "title": "Triage Required",
                "message": "No triage decision was detected. The intelligence officer "
                           "must assign a triage decision before submission.",
                "priority": "medium",
            })

    elif target_unit == "operations":
        if "warrant_basis" not in extracted_fields:
            guidance.append({
                "type": "compliance",
                "title": "Legal Basis Not Specified",
                "message": "No warrant or legal basis was detected. Every operation "
                           "requires a documented legal authority (Firearms Act s.91, "
                           "DDA s.21, POCA s.14, etc.).",
                "priority": "high",
            })
        if "team_lead" not in extracted_fields:
            guidance.append({
                "type": "action",
                "title": "Team Lead Not Identified",
                "message": "Assign a team lead (badge number) for this operation.",
                "priority": "medium",
            })

    elif target_unit == "seizures":
        subtype = extracted_fields.get("_subtype", "firearms")
        if subtype == "firearms":
            if "serial_number" not in extracted_fields:
                guidance.append({
                    "type": "compliance",
                    "title": "Serial Number Missing",
                    "message": "No serial number was detected. If the serial is "
                               "obliterated, record as 'OBLITERATED' and ensure "
                               "IBIS/eTrace status reflects this.",
                    "priority": "high",
                })
            guidance.append({
                "type": "action",
                "title": "IBIS & eTrace Submission Required",
                "message": "Submit this firearm to IBIS (ballistic database) and "
                           "eTrace (ATF tracing) as soon as possible. Set initial "
                           "status to 'Submitted - Pending'.",
                "priority": "high",
            })
            guidance.append({
                "type": "compliance",
                "title": "Evidence Chain of Custody",
                "message": "Ensure a corresponding Chain of Custody record is "
                           "created in the Forensics unit. The exhibit must be "
                           "sealed, tagged, and photographed in situ.",
                "priority": "high",
            })
        else:
            guidance.append({
                "type": "compliance",
                "title": "Lab Submission Required",
                "message": "Submit a sample to IFSLM for chemical analysis. A "
                           "forensic certificate is required for DPP file.",
                "priority": "high",
            })
            if "est_street_value" not in extracted_fields:
                guidance.append({
                    "type": "action",
                    "title": "Street Value Not Estimated",
                    "message": "Provide an estimated street value. This is required "
                               "for the prosecution file and POCA assessment.",
                    "priority": "medium",
                })

    elif target_unit == "arrests":
        guidance.append({
            "type": "compliance",
            "title": "48-Hour Rule (Constabulary Force Act s.15)",
            "message": "The suspect must be brought before a court within 48 hours "
                       "of arrest. Track the charge deadline carefully.",
            "priority": "high",
        })
        if "offence_1" not in extracted_fields:
            guidance.append({
                "type": "action",
                "title": "Offence Not Identified",
                "message": "The bot could not identify a specific offence from the "
                           "report. Select the appropriate offence and law section.",
                "priority": "high",
            })
        guidance.append({
            "type": "compliance",
            "title": "Suspect Rights",
            "message": "Ensure the suspect has been cautioned (Judges' Rules), "
                       "constitutional rights advised (Chapter III), and access "
                       "to attorney has been facilitated.",
            "priority": "high",
        })

    elif target_unit == "forensics":
        guidance.append({
            "type": "compliance",
            "title": "Chain of Custody Protocol",
            "message": "Ensure the exhibit is sealed in a tamper-evident bag, "
                       "tagged with a unique exhibit number, and all transfers "
                       "are recorded.",
            "priority": "high",
        })

    elif target_unit == "registry":
        guidance.append({
            "type": "compliance",
            "title": "SOP Compliance Checklist",
            "message": "An SOP compliance checklist should be created alongside "
                       "this case registration. Review the checklist items for "
                       "initial response, suspect processing, and evidence.",
            "priority": "medium",
        })
        guidance.append({
            "type": "info",
            "title": "DPP File Timeline",
            "message": "Begin assembling the DPP file package. Non-uniformed "
                       "officers must submit to DCR within 72 hours. Ensure "
                       "all forensic certificates are obtained.",
            "priority": "medium",
        })

    # --- General guidance ---

    if not text.strip():
        guidance.append({
            "type": "warning",
            "title": "No Text Extracted",
            "message": "The bot could not extract text from this file. This may "
                       "be a scanned image. Please enter the information manually "
                       "in the fields below.",
            "priority": "high",
        })

    return guidance


# ---------------------------------------------------------------------------
# Main processing entry point
# ---------------------------------------------------------------------------

def process_report(filepath, user_unit=None):
    """Process an uploaded report file.

    Args:
        filepath: Path to the uploaded file.
        user_unit: The unit the user belongs to (if single-unit user).

    Returns:
        dict with keys:
            text: Raw extracted text
            target_unit: Best-match unit key
            unit_scores: List of (unit, score%) tuples
            fields: Dict of extracted field values
            guidance: List of guidance dicts
            subtype: Optional subtype (e.g. 'firearms' or 'narcotics')
    """
    text = extract_text(filepath)

    if user_unit:
        target_unit = user_unit
        unit_scores = [(user_unit, 100)]
    else:
        unit_scores = detect_units(text)
        target_unit = unit_scores[0][0] if unit_scores else "intel"

    fields = extract_fields(text, target_unit)
    guidance = generate_guidance(text, target_unit, fields)

    subtype = fields.pop("_subtype", None)
    # Remove internal fields from the output
    internal = [k for k in fields if k.startswith("_")]
    meta = {k: fields.pop(k) for k in internal}

    return {
        "text": text,
        "target_unit": target_unit,
        "unit_scores": unit_scores,
        "fields": fields,
        "meta": meta,
        "guidance": guidance,
        "subtype": subtype,
    }
