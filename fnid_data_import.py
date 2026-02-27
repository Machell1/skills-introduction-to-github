#!/usr/bin/env python3
"""
FNID Data Import Tool v1.0

Extracts data from Microsoft Word documents (.docx) and Excel spreadsheets
(.xlsx) and pushes records into the appropriate FNID unit portal workbooks.

Supports:
  - Structured Word templates (tables with labelled fields)
  - Free-form Word reports (keyword-based extraction)
  - Excel spreadsheets (column mapping)
  - Auto-fill: parses source documents and pre-populates portal fields
  - Manual entry: interactive CLI for officer data entry

Usage:
  python fnid_data_import.py --source report.docx --unit intel
  python fnid_data_import.py --source data.xlsx --unit seizures --sheet Firearms
  python fnid_data_import.py --manual --unit arrests
  python fnid_data_import.py --list-units

Applicable legislation references extracted automatically:
  - Firearms Act 2022 sections
  - Dangerous Drugs Act sections
  - POCA 2007 references
  - Bail Act 2023 references
"""

import argparse
import os
import re
import sys
from datetime import datetime

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

PORTAL_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "portals")

# Unit portal file mapping
UNIT_PORTALS = {
    "intel": {
        "file": "FNID_Intel_Portal.xlsx",
        "sheet": "Intel_Log",
        "id_prefix": "INT",
    },
    "operations": {
        "file": "FNID_Operations_Portal.xlsx",
        "sheet": "Op_Register",
        "id_prefix": "OP",
    },
    "seizures-firearms": {
        "file": "FNID_Seizures_Portal.xlsx",
        "sheet": "Firearm_Seizures",
        "id_prefix": "FS",
    },
    "seizures-narcotics": {
        "file": "FNID_Seizures_Portal.xlsx",
        "sheet": "Narcotics_Seizures",
        "id_prefix": "NS",
    },
    "seizures-other": {
        "file": "FNID_Seizures_Portal.xlsx",
        "sheet": "Other_Seizures",
        "id_prefix": "OS",
    },
    "arrests": {
        "file": "FNID_ArrestsCourt_Portal.xlsx",
        "sheet": "Arrest_Register",
        "id_prefix": "AR",
    },
    "forensics-custody": {
        "file": "FNID_Forensics_Portal.xlsx",
        "sheet": "Chain_of_Custody",
        "id_prefix": "EXH",
    },
    "forensics-lab": {
        "file": "FNID_Forensics_Portal.xlsx",
        "sheet": "Lab_Tracking",
        "id_prefix": "LAB",
    },
    "registry": {
        "file": "FNID_Registry_Portal.xlsx",
        "sheet": "Case_Registry",
        "id_prefix": "CR",
    },
    "registry-dpp": {
        "file": "FNID_Registry_Portal.xlsx",
        "sheet": "DPP_Pipeline",
        "id_prefix": "DPP",
    },
    "registry-sop": {
        "file": "FNID_Registry_Portal.xlsx",
        "sheet": "SOP_Checklist",
        "id_prefix": "SOP",
    },
    "registry-witness": {
        "file": "FNID_Registry_Portal.xlsx",
        "sheet": "Witness_Statements",
        "id_prefix": "WIT",
    },
    "registry-disclosure": {
        "file": "FNID_Registry_Portal.xlsx",
        "sheet": "Disclosure_Log",
        "id_prefix": "DISC",
    },
}

# Keyword patterns for auto-detection of content type from Word docs
KEYWORD_PATTERNS = {
    "intel": [
        r"intelligence\s+report", r"intel\s+report", r"informant",
        r"crime\s*stop", r"tip\s+received", r"source\s+information",
        r"surveillance\s+report",
    ],
    "operations": [
        r"operation\s+order", r"op\s+order", r"search\s+warrant",
        r"warrant\s+execution", r"raid\s+report", r"operation\s+report",
        r"joint\s+operation", r"op\s+\w+\s+report",
    ],
    "seizures-firearms": [
        r"firearm\s+seizure", r"gun\s+seized", r"pistol",
        r"rifle", r"shotgun", r"ammunition\s+seized",
        r"weapon\s+recovered", r"ballistic",
    ],
    "seizures-narcotics": [
        r"drug\s+seizure", r"narcotic", r"cocaine", r"ganja",
        r"cannabis", r"heroin", r"controlled\s+substance",
    ],
    "arrests": [
        r"arrest\s+report", r"suspect\s+arrested", r"charged?\s+with",
        r"caution(?:ed)?", r"rights\s+advis", r"48.?hour",
        r"detain", r"bail\s+(?:granted|denied)",
    ],
    "forensics-custody": [
        r"exhibit", r"chain\s+of\s+custody", r"evidence\s+tag",
        r"sealed", r"armoury", r"storage",
    ],
    "registry": [
        r"case\s+file", r"case\s+registry", r"dpp\s+submission",
        r"prosecution\s+file", r"case\s+management",
        r"sop\s+compliance", r"disclosure",
    ],
}

# Law/section pattern extraction
LAW_PATTERNS = [
    (r"s(?:ection)?\.?\s*(\d+)\s*.*?firearms\s+act",
     "Firearms Act 2022 s.{}"),
    (r"s(?:ection)?\.?\s*(\d+)\s*.*?dangerous\s+drugs?\s+act",
     "DDA s.{}"),
    (r"s(?:ection)?\.?\s*(\d+)\s*.*?poca|proceeds\s+of\s+crime",
     "POCA 2007 s.{}"),
    (r"s(?:ection)?\.?\s*(\d+)\s*.*?bail\s+act",
     "Bail Act 2023 s.{}"),
    (r"s(?:ection)?\.?\s*(\d+)\s*.*?gun\s+court\s+act",
     "Gun Court Act s.{}"),
]

# Parish detection
PARISHES = [
    "Manchester", "St. Elizabeth", "Clarendon", "Kingston",
    "St. Andrew", "St. Thomas", "Portland", "St. Mary",
    "St. Ann", "Trelawny", "St. James", "Hanover",
    "Westmoreland", "St. Catherine",
]


# =================================================================
#  WORD DOCUMENT EXTRACTION
# =================================================================
def extract_from_word(filepath):
    """Extract structured and free-form data from a Word document.

    Returns a dict with:
      - 'tables': list of dicts (one per Word table, key=first-col, val=second-col)
      - 'text': full text as string
      - 'detected_unit': best-guess unit type
      - 'fields': extracted field values
      - 'laws': detected law references
      - 'parish': detected parish
    """
    if DocxDocument is None:
        print("ERROR: python-docx not installed. Run: pip install python-docx")
        sys.exit(1)

    doc = DocxDocument(filepath)
    result = {
        "tables": [],
        "text": "",
        "detected_unit": None,
        "fields": {},
        "laws": [],
        "parish": None,
    }

    # Extract all text
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    result["text"] = full_text

    # Extract tables (structured data)
    for table in doc.tables:
        table_data = {}
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0]:
                # Key-value pair (label in first column, value in second)
                key = cells[0].rstrip(":").strip()
                val = cells[1].strip()
                table_data[key.lower()] = val
            elif len(cells) >= 1:
                # Single-column or multi-column data
                for i, c in enumerate(cells):
                    if c:
                        table_data[f"col{i}_{len(table_data)}"] = c
        if table_data:
            result["tables"].append(table_data)

    # Auto-detect unit type from content
    text_lower = full_text.lower()
    scores = {}
    for unit, patterns in KEYWORD_PATTERNS.items():
        score = sum(1 for p in patterns if re.search(p, text_lower))
        if score > 0:
            scores[unit] = score

    if scores:
        result["detected_unit"] = max(scores, key=scores.get)

    # Extract law references
    for pattern, template in LAW_PATTERNS:
        matches = re.findall(pattern, text_lower)
        for m in matches:
            result["laws"].append(template.format(m))

    # Detect parish
    for parish in PARISHES:
        if parish.lower() in text_lower:
            result["parish"] = parish
            break

    # Extract common fields from text using patterns
    fields = result["fields"]

    # Date patterns
    date_match = re.search(
        r"(?:date|dated?)[\s:]*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})", text_lower)
    if date_match:
        fields["date"] = date_match.group(1)

    # Name patterns
    name_match = re.search(
        r"(?:suspect|accused|target|person)[\s:]*([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)",
        full_text)
    if name_match:
        fields["suspect_name"] = name_match.group(1)

    # Location patterns
    loc_match = re.search(
        r"(?:location|address|premises|at)\s*[:\-]?\s*(.+?)(?:\.|$)",
        full_text, re.MULTILINE)
    if loc_match:
        fields["location"] = loc_match.group(1).strip()[:60]

    # Officer patterns
    officer_match = re.search(
        r"(?:officer|investigating|oic|arresting|seized by)[\s:]*([A-Z][a-z]+\.?\s+[A-Z][a-z]+)",
        full_text)
    if officer_match:
        fields["officer"] = officer_match.group(1)

    # Also check table data for common field names
    for table_data in result["tables"]:
        for key, val in table_data.items():
            if any(k in key for k in ["date", "time"]):
                fields.setdefault("date", val)
            elif any(k in key for k in ["name", "suspect", "accused"]):
                fields.setdefault("suspect_name", val)
            elif any(k in key for k in ["parish", "division"]):
                fields.setdefault("parish", val)
            elif any(k in key for k in ["location", "address"]):
                fields.setdefault("location", val)
            elif any(k in key for k in ["officer", "oic", "rank"]):
                fields.setdefault("officer", val)
            elif any(k in key for k in ["offence", "charge", "section"]):
                fields.setdefault("offence", val)
            elif any(k in key for k in ["case", "reference", "ref"]):
                fields.setdefault("case_ref", val)
            elif any(k in key for k in ["firearm", "weapon", "gun"]):
                fields.setdefault("firearm_type", val)
            elif any(k in key for k in ["drug", "narcotic", "substance"]):
                fields.setdefault("drug_type", val)
            elif any(k in key for k in ["quantity", "amount", "weight"]):
                fields.setdefault("quantity", val)

    if result["parish"] and "parish" not in fields:
        fields["parish"] = result["parish"]

    return result


# =================================================================
#  EXCEL SPREADSHEET EXTRACTION
# =================================================================
def extract_from_excel(filepath, sheet_name=None):
    """Extract data rows from an Excel spreadsheet.

    Returns a list of dicts, one per row, keyed by header names.
    """
    wb = load_workbook(filepath, data_only=True)

    if sheet_name and sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        ws = wb.active
        if sheet_name:
            print(f"  Warning: Sheet '{sheet_name}' not found. Using '{ws.title}'.")

    # Read headers from row 1
    headers = []
    for col in range(1, ws.max_column + 1):
        val = ws.cell(row=1, column=col).value
        headers.append(str(val).strip() if val else f"Column{col}")

    # Read data rows
    rows = []
    for r in range(2, ws.max_row + 1):
        row_data = {}
        has_data = False
        for c, header in enumerate(headers, 1):
            val = ws.cell(row=r, column=c).value
            if val is not None:
                has_data = True
            row_data[header.lower()] = val
        if has_data:
            rows.append(row_data)

    wb.close()
    return rows


# =================================================================
#  FIELD MAPPING - Maps extracted fields to portal columns
# =================================================================
FIELD_MAPS = {
    "intel": {
        "date": "B",     # DateReceived
        "source": "D",   # Source
        "parish": "M",   # Parish
        "location": "L",  # TargetLocation
        "suspect_name": "K",  # TargetPerson
        "officer": "V",  # CreatedBy
    },
    "operations": {
        "date": "C",     # OpDate
        "op_type": "D",  # OpType
        "parish": "H",   # Parish
        "location": "I",  # TargetLocation
        "suspect_name": "J",  # TargetPerson
        "officer": "K",  # TeamLead
    },
    "seizures-firearms": {
        "date": "B",     # SeizureDate
        "parish": "D",   # Parish
        "location": "E",  # Location
        "firearm_type": "F",  # FirearmType
        "officer": "R",  # SeizedBy
    },
    "seizures-narcotics": {
        "date": "B",     # SeizureDate
        "parish": "D",   # Parish
        "location": "E",  # Location
        "drug_type": "F",  # DrugType
        "quantity": "G",  # Quantity
        "officer": "P",  # SeizedBy
    },
    "arrests": {
        "date": "B",      # ArrestDate
        "suspect_name": "E",  # SuspectName
        "parish": "I",    # Parish
        "location": "J",  # ArrestLocation
        "officer": "K",   # ArrestingOfficer
        "offence": "M",   # Offence1
    },
    "registry": {
        "date": "B",       # RegistrationDate
        "officer": "D",    # OIC_Name
        "parish": "G",     # Parish
        "offence": "I",    # OffenceDescription
        "suspect_name": "K",  # SuspectName
    },
}


# =================================================================
#  IMPORT TO PORTAL
# =================================================================
def get_next_id(ws, id_prefix, year=None):
    """Find the next available ID in column A."""
    if year is None:
        year = datetime.now().year
    max_num = 0
    for r in range(2, ws.max_row + 1):
        val = ws.cell(row=r, column=1).value
        if val and str(val).startswith(f"{id_prefix}-{year}-"):
            try:
                num = int(str(val).split("-")[-1])
                max_num = max(max_num, num)
            except ValueError:
                pass
    return f"{id_prefix}-{year}-{str(max_num + 1).zfill(3)}"


def find_next_empty_row(ws):
    """Find the next empty row in the worksheet (within table range)."""
    for r in range(2, ws.max_row + 1):
        if ws.cell(row=r, column=1).value is None:
            return r
    return ws.max_row + 1


def import_to_portal(unit_key, fields, officer_name=None):
    """Import extracted fields into the appropriate portal workbook.

    Returns (success, message) tuple.
    """
    if unit_key not in UNIT_PORTALS:
        return False, f"Unknown unit: {unit_key}"

    config = UNIT_PORTALS[unit_key]
    portal_path = os.path.join(PORTAL_DIR, config["file"])

    if not os.path.exists(portal_path):
        return False, f"Portal not found: {portal_path}. Run fnid_portal_system.py first."

    wb = load_workbook(portal_path)
    ws = wb[config["sheet"]]

    # Find next empty row
    row = find_next_empty_row(ws)

    # Generate ID
    new_id = get_next_id(ws, config["id_prefix"])
    ws.cell(row=row, column=1, value=new_id)

    # Map fields to columns
    base_unit = unit_key.split("-")[0] if "-" in unit_key else unit_key
    field_map = FIELD_MAPS.get(base_unit, FIELD_MAPS.get(unit_key, {}))

    mapped_count = 0
    for field_name, col_letter in field_map.items():
        if field_name in fields and fields[field_name]:
            col_idx = ord(col_letter) - ord("A") + 1
            ws.cell(row=row, column=col_idx, value=fields[field_name])
            mapped_count += 1

    # Set submission tracking
    # Find the RecordStatus column (last few columns)
    for c in range(ws.max_column, max(1, ws.max_column - 10), -1):
        if ws.cell(row=1, column=c).value == "RecordStatus":
            ws.cell(row=row, column=c, value="Draft")
            if officer_name:
                ws.cell(row=row, column=c + 1, value=officer_name)
            ws.cell(row=row, column=c + 2,
                    value=datetime.now().strftime("%Y-%m-%d"))
            break

    wb.save(portal_path)
    wb.close()
    return True, f"Record {new_id} added to {config['sheet']} ({mapped_count} fields mapped) in {config['file']}"


# =================================================================
#  MANUAL ENTRY
# =================================================================
def manual_entry(unit_key):
    """Interactive CLI for manual data entry."""
    if unit_key not in UNIT_PORTALS:
        print(f"Unknown unit: {unit_key}")
        return

    config = UNIT_PORTALS[unit_key]
    portal_path = os.path.join(PORTAL_DIR, config["file"])

    if not os.path.exists(portal_path):
        print(f"Portal not found: {portal_path}")
        return

    wb = load_workbook(portal_path)
    ws = wb[config["sheet"]]

    # Read column headers
    headers = []
    for c in range(1, ws.max_column + 1):
        val = ws.cell(row=1, column=c).value
        if val and val not in ["RecordStatus", "SubmittedBy", "SubmittedDate",
                               "LastEditedBy", "LastEditedDate"]:
            headers.append((c, str(val)))
        else:
            break

    print(f"\n--- Manual Entry: {config['sheet']} ---")
    print(f"Portal: {config['file']}")
    print(f"Enter values for each field (press Enter to skip):\n")

    row = find_next_empty_row(ws)
    new_id = get_next_id(ws, config["id_prefix"])
    ws.cell(row=row, column=1, value=new_id)
    print(f"  [Auto] {headers[0][1]}: {new_id}")

    filled = 0
    for col_idx, header in headers[1:]:
        try:
            val = input(f"  {header}: ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n  Entry cancelled.")
            wb.close()
            return
        if val:
            ws.cell(row=row, column=col_idx, value=val)
            filled += 1

    # Set status
    for c in range(ws.max_column, max(1, ws.max_column - 10), -1):
        if ws.cell(row=1, column=c).value == "RecordStatus":
            ws.cell(row=row, column=c, value="Draft")
            ws.cell(row=row, column=c + 2,
                    value=datetime.now().strftime("%Y-%m-%d"))
            break

    wb.save(portal_path)
    wb.close()
    print(f"\n  Record {new_id} saved ({filled} fields filled).")
    print(f"  Open {config['file']} to review and set status to 'Submitted'.\n")


# =================================================================
#  MAIN CLI
# =================================================================
def main():
    parser = argparse.ArgumentParser(
        description="FNID Data Import Tool - Extract from Word/Excel to portals",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s --source report.docx --unit intel
  %(prog)s --source report.docx --auto-detect
  %(prog)s --source data.xlsx --unit seizures-firearms --sheet Sheet1
  %(prog)s --manual --unit arrests
  %(prog)s --list-units
        """)

    parser.add_argument("--source", "-s",
                        help="Source file path (.docx or .xlsx)")
    parser.add_argument("--unit", "-u",
                        help="Target unit portal key (use --list-units to see options)")
    parser.add_argument("--sheet",
                        help="Sheet name for Excel source files")
    parser.add_argument("--auto-detect", action="store_true",
                        help="Auto-detect unit from document content")
    parser.add_argument("--manual", action="store_true",
                        help="Interactive manual entry mode")
    parser.add_argument("--officer", "-o",
                        help="Officer name for submission tracking")
    parser.add_argument("--list-units", action="store_true",
                        help="List available unit portal keys")

    args = parser.parse_args()

    if args.list_units:
        print("\nAvailable unit portal keys:")
        print("-" * 60)
        for key, config in UNIT_PORTALS.items():
            print(f"  {key:25s}  -> {config['file']} / {config['sheet']}")
        print()
        return

    if args.manual:
        if not args.unit:
            print("ERROR: --unit required for manual entry. Use --list-units.")
            sys.exit(1)
        manual_entry(args.unit)
        return

    if not args.source:
        parser.print_help()
        sys.exit(1)

    if not os.path.exists(args.source):
        print(f"ERROR: Source file not found: {args.source}")
        sys.exit(1)

    ext = os.path.splitext(args.source)[1].lower()

    if ext == ".docx":
        print(f"\nExtracting from Word document: {args.source}")
        result = extract_from_word(args.source)

        print(f"  Text length: {len(result['text'])} chars")
        print(f"  Tables found: {len(result['tables'])}")
        print(f"  Detected unit: {result['detected_unit'] or 'Unknown'}")
        print(f"  Parish: {result['parish'] or 'Not detected'}")
        if result["laws"]:
            print(f"  Law references: {', '.join(result['laws'])}")
        print(f"  Fields extracted: {len(result['fields'])}")
        for k, v in result["fields"].items():
            print(f"    {k}: {v}")

        # Determine target unit
        unit = args.unit
        if not unit and args.auto_detect:
            unit = result["detected_unit"]
        if not unit:
            print("\n  Could not determine target unit.")
            print("  Use --unit <key> or --auto-detect. See --list-units.")
            sys.exit(1)

        # Import
        success, msg = import_to_portal(unit, result["fields"], args.officer)
        if success:
            print(f"\n  SUCCESS: {msg}")
        else:
            print(f"\n  ERROR: {msg}")

    elif ext == ".xlsx":
        print(f"\nExtracting from Excel spreadsheet: {args.source}")
        rows = extract_from_excel(args.source, args.sheet)
        print(f"  Rows extracted: {len(rows)}")

        if not args.unit:
            print("  ERROR: --unit required for Excel import.")
            sys.exit(1)

        imported = 0
        for row_data in rows:
            success, msg = import_to_portal(args.unit, row_data, args.officer)
            if success:
                imported += 1
                print(f"  {msg}")
            else:
                print(f"  ERROR: {msg}")

        print(f"\n  Imported {imported}/{len(rows)} records to {args.unit}.")

    else:
        print(f"ERROR: Unsupported file type: {ext}")
        print("  Supported: .docx (Word), .xlsx (Excel)")
        sys.exit(1)


if __name__ == "__main__":
    main()
